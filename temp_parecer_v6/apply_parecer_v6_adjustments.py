from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\temp\PDS")
SRC = Path(r"C:\Users\android\Downloads\AUTODOWNLOAD (1).docx")
OUT = ROOT / "Documents" / "AUTODOWNLOAD_ajustado_parecer_v6.docx"
IMG_DIR = ROOT / "temp_parecer_v6" / "generated_images"
LOG = ROOT / "temp_parecer_v6" / "changes.txt"

TEST_COMMAND = r"dotnet test .\backend\tests\AutoDownload.Tests\AutoDownload.Tests.csproj --no-restore"
TEST_TRANSCRIPT = (
    "Execução de teste para C:\\temp\\PDS\\backend\\tests\\AutoDownload.Tests\\bin\\Debug\\net10.0\\AutoDownload.Tests.dll (.NETCoreApp,Version=v10.0)\n"
    "Versão do VSTest 18.0.1 (x64)\n"
    "Aprovado! - Com falha: 0, Aprovado: 20, Ignorado: 0, Total: 20, Duração: 182 ms - AutoDownload.Tests.dll (net10.0)"
)


def font(size: int, bold: bool = False):
    path = Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf")
    return ImageFont.truetype(str(path), size=size) if path.exists() else ImageFont.load_default()


def wrap(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> str:
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for word in raw.split():
            candidate = word if not current else f"{current} {word}"
            if draw.textlength(candidate, font=fnt) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        lines.append(current)
    return "\n".join(lines)


def draw_entity(draw: ImageDraw.ImageDraw, title: str, fields: list[str], box: tuple[int, int, int, int]) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=14, fill="#f8fafc", outline="#1d4ed8", width=3)
    draw.rectangle((x1, y1, x2, y1 + 44), fill="#1d4ed8")
    draw.text((x1 + 14, y1 + 9), title, font=font(22, True), fill="#ffffff")
    y = y1 + 58
    for field in fields:
        draw.text((x1 + 14, y), field, font=font(18), fill="#0f172a")
        y += 23


def draw_line(draw: ImageDraw.ImageDraw, a: tuple[int, int], b: tuple[int, int]) -> None:
    draw.line((a, b), fill="#475569", width=3)


def make_data_model(path: Path) -> None:
    img = Image.new("RGB", (1800, 1200), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((50, 35), "Modelo de dados do AutoDownload", font=font(36, True), fill="#0f172a")
    entities = {
        "users": ((70, 145, 460, 335), ["Id", "Name", "Email", "PasswordHash", "CreatedAt", "UpdatedAt"]),
        "operators": ((1280, 145, 1720, 335), ["Id", "Code", "Name", "ServiceType", "PortalBaseUrl", "IsActive"]),
        "accounts": (
            (560, 135, 1180, 485),
            [
                "Id",
                "UserId",
                "OperatorId",
                "PortalLogin",
                "CustomerIdentifier",
                "Status",
                "LastRunAt",
                "NextRunAt",
                "IsScheduleEnabled",
                "ScheduleDayOfMonth",
                "ScheduleTime",
            ],
        ),
        "bills": ((70, 650, 520, 950), ["Id", "UserId", "AccountId", "OperatorId", "Reference", "DueDate", "Amount", "FileName", "StoragePath", "DownloadedAt", "Status"]),
        "automation_runs": ((650, 650, 1120, 950), ["Id", "UserId", "AccountId", "OperatorId", "StartedAt", "FinishedAt", "Status", "Message", "FileName"]),
        "notifications": ((1260, 650, 1720, 870), ["Id", "UserId", "Type", "Text", "CreatedAt", "ReadAt"]),
    }
    for name, (box, fields) in entities.items():
        draw_entity(draw, name, fields, box)
    centers = {name: ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2) for name, (box, _) in entities.items()}
    for a, b, la, lb in [
        ("users", "accounts", "1", "N"),
        ("operators", "accounts", "1", "N"),
        ("accounts", "bills", "1", "N"),
        ("accounts", "automation_runs", "1", "N"),
        ("users", "notifications", "1", "N"),
    ]:
        draw_line(draw, centers[a], centers[b])
        draw.text((centers[a][0] + 8, centers[a][1] + 6), la, font=font(20, True), fill="#0f172a")
        draw.text((centers[b][0] + 8, centers[b][1] + 6), lb, font=font(20, True), fill="#0f172a")
    draw.text(
        (70, 1090),
        "A entidade accounts inclui os campos de agendamento mensal: IsScheduleEnabled, ScheduleDayOfMonth e ScheduleTime.",
        font=font(22, True),
        fill="#1e3a8a",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def set_text(paragraph: Paragraph, text: str, bold: bool | None = None, size: int | None = None, align=None, font_name: str = "Arial") -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
    if bold is not None:
        run.font.bold = bold
    if size:
        run.font.size = Pt(size)
    if align is not None:
        paragraph.alignment = align


def style_runs(paragraph: Paragraph, size: int = 12, font_name: str = "Arial") -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)


def insert_after(paragraph: Paragraph, text: str, size: int = 12, font_name: str = "Arial") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_text(new_para, text, size=size, font_name=font_name)
    return new_para


def keep_with_next(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def replace_in_runs(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        if run._element.xpath(".//pic:pic"):
            continue
        text = run.text
        for old, new in replacements.items():
            text = text.replace(old, new)
        run.text = text


def image_caption_same_paragraph(paragraph: Paragraph, caption: str) -> None:
    done = False
    for run in paragraph.runs:
        if run._element.xpath(".//pic:pic"):
            continue
        if not done:
            run.text = f"\n{caption}\nFonte: Elaborado pelo autor (2026)."
            done = True
        else:
            run.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_image(paragraph: Paragraph, image_path: Path, width_cm: float = 16.0) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    paragraph.add_run("\nFonte: Elaborado pelo autor (2026).")


def remove_row(table, index: int) -> None:
    table._tbl.remove(table.rows[index]._tr)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                style_runs(paragraph, size=9)
                if row_i == 0:
                    for run in paragraph.runs:
                        run.font.bold = True


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    data_model = IMG_DIR / "modelo_dados_v6.png"
    make_data_model(data_model)

    doc = Document(str(SRC))

    # 1. Título.
    set_text(
        doc.paragraphs[0],
        "AUTODOWNLOAD: UMA APLICAÇÃO WEB PARA AUTOMAÇÃO E GESTÃO DE BOLETOS RECORRENTES",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 2. Resumo com 150+ palavras.
    resumo = doc.paragraphs[5]
    set_text(
        resumo,
        "O presente Trabalho de Conclusão de Curso apresenta o desenvolvimento do AutoDownload, uma aplicação web voltada à automação do acesso a portais de operadoras de internet, telefonia e telecomunicações, bem como à gestão de boletos recorrentes. O objetivo do projeto é centralizar o cadastro de contas, executar automações de download, registrar boletos, histórico e notificações, reduzindo atividades manuais repetitivas. A metodologia combinou pesquisa aplicada, desenvolvimento tecnológico, levantamento de requisitos inspirado em Design Thinking, prototipação de interface e validação técnica da aplicação. O frontend foi desenvolvido com Next.js e React, enquanto o backend foi implementado em C# com .NET 10, Entity Framework Core, PostgreSQL, autenticação JWT e estratégias de automação com Selenium WebDriver. A validação incluiu uso manual do sistema, verificação da API e execução de testes automatizados no backend, com 20 testes aprovados e nenhuma falha. Como limitação, destaca-se a dependência de portais externos, que podem alterar fluxos, bloquear automações ou exigir manutenção contínua.",
        size=12,
    )

    # 3, 4, 5. Concordância, .NET, palavras-chave, escopo.
    replacements = {
        "Tecnologias Digitais Aplicada a Educação": "Tecnologias Digitais Aplicadas à Educação",
        "Tecnologias Digitais Aplicadas a Educação": "Tecnologias Digitais Aplicadas à Educação",
        "Palavras-chave: Automação web. Boletos recorrentes. NET. Next.js. PostgreSQL. Selenium.": "Palavras-chave: Automação web. Boletos recorrentes. Aplicação web. .NET. Selenium.",
        "internet, energia, água, telefone ou outros serviços recorrentes": "internet, telefonia e serviços de telecomunicações",
        "serviços financeiros, de telecomunicações e de utilidades públicas": "serviços financeiros, de internet, telefonia e telecomunicações",
        "REFERĘNCIAS": "REFERÊNCIAS",
        "Proteçăo": "Proteção",
        "Presidęncia": "Presidência",
        "Săo": "São",
    }
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph, replacements)
    set_text(doc.paragraphs[7], "Palavras-chave: Automação web. Boletos recorrentes. Aplicação web. .NET. Selenium.", size=12)

    # 6. Figura 1: legenda padronizada e grudada na imagem.
    set_text(doc.paragraphs[54], "Figura 1 - Diagrama de casos de uso do AutoDownload.", bold=True, size=12)
    doc.paragraphs[54].paragraph_format.page_break_before = True
    keep_with_next(doc.paragraphs[54])
    keep_with_next(doc.paragraphs[55])

    # 7. Numeração contínua das figuras do corpo e apêndice D.
    set_text(doc.paragraphs[64], "Figura 2 - Arquitetura geral do AutoDownload.", bold=True, size=12)
    set_text(doc.paragraphs[72], "Figura 3 - Modelo de dados do AutoDownload.", bold=True, size=12)
    replace_image(doc.paragraphs[73], data_model)
    image_caption_same_paragraph(doc.paragraphs[77], "Figura 4 - Tela de login.")
    image_caption_same_paragraph(doc.paragraphs[78], "Figura 5 - Tela inicial da aplicação.")
    image_caption_same_paragraph(doc.paragraphs[79], "Figura 6 - Tela para adicionar uma conta.")
    set_text(doc.paragraphs[81], "Figura 7 - Menu lateral com opções e modo claro/escuro.\nFonte: Elaborado pelo autor (2026).", bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    appendix_caption_replacements = {
        "Figura D1: Tela de login.": "Figura D1 - Tela de login.",
        "Figura D2: Tela principal.": "Figura D2 - Tela principal.",
        "Figura D3: Tela de contas cadastradas e cadastro de contas.": "Figura D3 - Tela de contas cadastradas e cadastro de contas.",
        "Figura D4: Tela de notificações.": "Figura D4 - Tela de notificações.",
        "Figura D5: Tela de históricos de downloads.": "Figura D5 - Tela de históricos de downloads.",
        "Figura D6: Tela de contas com uma conta de demonstração.": "Figura D6 - Tela de contas com uma conta de demonstração.",
        "Figura D7: Tela de agendamento mensal para download automático do boleto.": "Figura D7 - Tela de agendamento mensal para download automático do boleto.",
        "Figura D9: Tela de boletos baixados, para poderem ser baixados novamente.": "Figura D8 - Tela de boletos baixados, para poderem ser baixados novamente.",
        "Figura D8: Menu lateral com opções e modo claro/escuro.": "Figura D9 - Menu lateral com opções e modo claro/escuro.",
        "Figura D9: Aparência da aplicação no modo escuro.": "Figura D10 - Aparência da aplicação no modo escuro.",
    }
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        for old, new in appendix_caption_replacements.items():
            if stripped.startswith(old):
                set_text(paragraph, new, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 9 e 10. Quadro 6 sem linhas contraditórias.
    validation = doc.tables[5]
    for idx in range(len(validation.rows) - 1, 0, -1):
        cells = [c.text.strip() for c in validation.rows[idx].cells]
        if cells[0] == "Automação real" or (cells[0] == "Agendamento mensal" and cells[-1] == "Aprovado"):
            remove_row(validation, idx)
    style_table(validation)

    # 11. Comando de testes formatado.
    set_text(
        doc.paragraphs[46],
        "A validação ocorreu por meio de build da solução, execução da suíte de testes unitários do backend com xUnit, verificação do endpoint de saúde da API, uso manual da aplicação e execuções controladas das automações. O comando utilizado para os testes automatizados do backend foi:",
        size=12,
    )
    result_para = insert_after(doc.paragraphs[46], "Resultado registrado em 25/06/2026: 20 testes aprovados, 0 falhas, 0 ignorados e duração de 182 ms.", size=12)
    insert_after(doc.paragraphs[46], TEST_COMMAND, size=10, font_name="Courier New")
    result_para.paragraph_format.space_after = Pt(6)

    # Appendix C test evidence.
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("Resultado registrado para a versão acadêmica: dotnet test"):
            set_text(paragraph, "Transcrição do teste automatizado do backend:", size=12)
            insert_after(paragraph, TEST_TRANSCRIPT, size=9, font_name="Courier New")
            insert_after(paragraph, TEST_COMMAND, size=10, font_name="Courier New")
            break

    # 12. RF/RNF padronizados.
    for table_index in (2, 3, 10):
        table = doc.tables[table_index]
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                text = re.sub(r"\bRF\s+0?(\d)\b", r"RF0\1", text)
                text = re.sub(r"\bRF\s+10\b", "RF10", text)
                text = text.replace("RF010", "RF10")
                text = re.sub(r"\bRNF\s+0?(\d)\b", r"RNF0\1", text)
                cell.text = text
        style_table(table)

    # 14. Registro da versão acadêmica claro por hash.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("APÊNDICE E - Registro da versão acadêmica"):
            set_text(paragraph, "APÊNDICE E - Registro da versão acadêmica: hash do commit e links do projeto.", bold=True, size=12)

    # 15. Referências online com [s.d.] quando não há ano editorial formal.
    ref_replacements = {
        "AUTOBOT. Repositório AutoBot. GitHub, 2026.": "AUTOBOT. Repositório AutoBot. GitHub, [s.d.].",
        "AUTODOWNLOAD. Repositório AutomacaoContas. GitHub, 2026.": "AUTODOWNLOAD. Repositório AutomacaoContas. GitHub, [s.d.].",
        "MICROSOFT. ASP.NET Core documentation. 2026a.": "MICROSOFT. ASP.NET Core documentation. [s.d.].",
        "MICROSOFT. Entity Framework Core documentation. 2026b.": "MICROSOFT. Entity Framework Core documentation. [s.d.].",
        "RMS TELECOM. Portal do assinante RMS Telecom. 2026.": "RMS TELECOM. Portal do assinante RMS Telecom. [s.d.].",
        "SELENIUM. Selenium WebDriver documentation. 2026.": "SELENIUM. Selenium WebDriver documentation. [s.d.].",
        "VERCEL. Next.js on Vercel. 2026.": "VERCEL. Next.js on Vercel. [s.d.].",
        "VERO INTERNET. Minha Vero. 2026.": "VERO INTERNET. Minha Vero. [s.d.].",
    }
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph, ref_replacements)

    # General formatting for touched captions/tables.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(("Figura", "Quadro", "APÊNDICE", "REFERÊNCIAS")):
            style_runs(paragraph, size=12)

    doc.save(str(OUT))
    LOG.parent.mkdir(parents=True, exist_ok=True)
    LOG.write_text(
        "\n".join(
            [
                f"Arquivo gerado: {OUT}",
                "Ajustes aplicados somente conforme Parecer V6:",
                "1. Título corrigido com espaço após dois-pontos.",
                "2. Resumo ampliado para mais de 150 palavras.",
                "3. Concordância da titulação do orientador corrigida.",
                "4. Palavras-chave reduzidas e .NET grafado corretamente.",
                "5. Escopo textual limitado a internet, telefonia e telecomunicações.",
                "6. Legenda da Figura 1 padronizada e mantida com a imagem.",
                "7. Numeração contínua das figuras do corpo e ajuste das figuras D no apêndice.",
                "8. Modelo de dados visual atualizado com campos de agendamento.",
                "9. Quadro 6 sem duplicidade contraditória.",
                "10. Apêndice C mantido coerente com validação parcial.",
                "11. Comando dotnet test formatado e transcrito.",
                "12. Códigos RF/RNF padronizados.",
                "14. Registro acadêmico identificado por hash.",
                "15. Referências online ajustadas com [s.d.] e acentos corrigidos.",
            ]
        ),
        encoding="utf-8",
    )
    print(OUT)


if __name__ == "__main__":
    main()
