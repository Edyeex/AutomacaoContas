from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\temp\PDS")
SRC = Path(r"C:\Users\android\Downloads\AUTODOWNLOAD.docx")
OUT = ROOT / "Documents" / "AUTODOWNLOAD_ajustado_parecer_v5.docx"
IMG_DIR = ROOT / "temp_parecer_v5" / "generated_images"
LOG = ROOT / "temp_parecer_v5" / "changes.txt"

TEST_RESULT = (
    "dotnet test .\\backend\\tests\\AutoDownload.Tests\\AutoDownload.Tests.csproj --no-restore; "
    "VSTest 18.0.1 x64; .NET 10; total: 20; aprovados: 20; falhas: 0; "
    "ignorados: 0; duração: 182 ms; execução em 25/06/2026."
)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=4)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> str:
    lines: list[str] = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if draw.textlength(candidate, font=fnt) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        lines.append(current)
    return "\n".join(lines)


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    title: bool = False,
    radius: int = 24,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    fnt = font(28 if title else 23, bold=title)
    wrapped = wrap_text(draw, text, fnt, x2 - x1 - 28)
    w, h = text_size(draw, wrapped, fnt)
    draw.multiline_text(
        (x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2),
        wrapped,
        font=fnt,
        fill="#0f172a",
        align="center",
        spacing=5,
    )


def draw_ellipse(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: str = "#eff6ff",
    outline: str = "#2563eb",
) -> None:
    x1, y1, x2, y2 = xy
    draw.ellipse(xy, fill=fill, outline=outline, width=3)
    fnt = font(21, bold=False)
    wrapped = wrap_text(draw, text, fnt, x2 - x1 - 30)
    w, h = text_size(draw, wrapped, fnt)
    draw.multiline_text(
        (x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2),
        wrapped,
        font=fnt,
        fill="#0f172a",
        align="center",
        spacing=4,
    )


def line(draw: ImageDraw.ImageDraw, a: tuple[int, int], b: tuple[int, int], fill: str = "#334155", width: int = 3) -> None:
    draw.line((a, b), fill=fill, width=width)


def dashed_line(draw: ImageDraw.ImageDraw, a: tuple[int, int], b: tuple[int, int], fill: str = "#64748b", width: int = 3) -> None:
    x1, y1 = a
    x2, y2 = b
    steps = 30
    for i in range(0, steps, 2):
        start = (x1 + (x2 - x1) * i / steps, y1 + (y2 - y1) * i / steps)
        end = (x1 + (x2 - x1) * (i + 1) / steps, y1 + (y2 - y1) * (i + 1) / steps)
        draw.line((start, end), fill=fill, width=width)


def actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    f = font(23, bold=True)
    draw.ellipse((x - 24, y, x + 24, y + 48), outline="#1e3a8a", width=4)
    draw.line((x, y + 48, x, y + 125), fill="#1e3a8a", width=4)
    draw.line((x - 55, y + 78, x + 55, y + 78), fill="#1e3a8a", width=4)
    draw.line((x, y + 125, x - 45, y + 185), fill="#1e3a8a", width=4)
    draw.line((x, y + 125, x + 45, y + 185), fill="#1e3a8a", width=4)
    wrapped = wrap_text(draw, label, f, 180)
    w, _ = text_size(draw, wrapped, f)
    draw.multiline_text((x - w / 2, y + 202), wrapped, font=f, fill="#0f172a", align="center")


def make_use_case_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "#ffffff")
    draw = ImageDraw.Draw(img)
    title_font = font(36, bold=True)
    draw.text((55, 35), "Diagrama de casos de uso do AutoDownload", font=title_font, fill="#0f172a")
    actor(draw, 140, 390, "Usuário")
    actor(draw, 1650, 390, "Portal da\noperadora")
    draw.rounded_rectangle((310, 115, 1490, 970), radius=28, outline="#1d4ed8", width=4, fill="#f8fafc")
    draw.text((350, 145), "Sistema AutoDownload", font=font(29, bold=True), fill="#1e3a8a")

    cases = {
        "Cadastrar usuário": (430, 225, 745, 310),
        "Fazer login": (430, 335, 745, 420),
        "Cadastrar conta": (430, 445, 745, 530),
        "Editar ou remover conta": (430, 555, 745, 640),
        "Executar automação manual": (820, 225, 1190, 310),
        "Configurar agendamento mensal": (820, 335, 1190, 420),
        "Consultar boletos": (820, 445, 1190, 530),
        "Ver histórico": (820, 555, 1190, 640),
        "Ver notificações": (820, 665, 1190, 750),
        "Autenticar no portal": (1210, 285, 1450, 360),
        "Consultar faturas": (1210, 405, 1450, 480),
        "Baixar boleto disponível": (1210, 525, 1450, 600),
    }
    for label, box in cases.items():
        draw_ellipse(draw, box, label)
    for box in list(cases.values())[:9]:
        line(draw, (205, 485), (box[0], (box[1] + box[3]) // 2), "#94a3b8", 2)
    for label in ["Autenticar no portal", "Consultar faturas", "Baixar boleto disponível"]:
        box = cases[label]
        line(draw, (box[2], (box[1] + box[3]) // 2), (1590, 485), "#94a3b8", 2)

    include_font = font(18)
    pairs = [
        ("Executar automação manual", "Autenticar no portal"),
        ("Executar automação manual", "Consultar faturas"),
        ("Executar automação manual", "Baixar boleto disponível"),
        ("Configurar agendamento mensal", "Executar automação manual"),
    ]
    for a, b in pairs:
        ax = cases[a]
        bx = cases[b]
        start = (ax[2], (ax[1] + ax[3]) // 2)
        end = (bx[0], (bx[1] + bx[3]) // 2)
        dashed_line(draw, start, end)
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 18
        draw.text((mx - 45, my), "<<include>>", font=include_font, fill="#475569")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1000), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((55, 35), "Arquitetura geral do AutoDownload", font=font(36, True), fill="#0f172a")
    boxes = [
        ((70, 410, 310, 545), "Usuário\nnavegador", "#e0f2fe", "#0369a1"),
        ((390, 360, 690, 595), "Frontend\nNext.js + React\nVercel", "#eff6ff", "#1d4ed8"),
        ((780, 320, 1100, 635), "API REST\n.NET 10\nJWT + controllers", "#eef2ff", "#4338ca"),
        ((1190, 165, 1640, 310), "PostgreSQL\nNeon\nusuários, contas,\nboletos e histórico", "#ecfdf5", "#047857"),
        ((1190, 420, 1640, 575), "Infraestrutura\nEF Core, criptografia,\nSelenium e strategies", "#f8fafc", "#475569"),
        ((1190, 680, 1640, 830), "Portais externos\nVero Internet / RMS Telecom", "#fff7ed", "#c2410c"),
    ]
    for box, label, fill, outline in boxes:
        draw_box(draw, box, label, fill, outline, title=False, radius=18)
    arrows = [
        ((310, 478), (390, 478), "HTTP"),
        ((690, 478), (780, 478), "API / JSON"),
        ((1100, 420), (1190, 245), "EF Core"),
        ((1100, 515), (1190, 500), "serviços"),
        ((1420, 575), (1420, 680), "Selenium"),
    ]
    for start, end, label in arrows:
        line(draw, start, end, "#1e293b", 4)
        lx = (start[0] + end[0]) // 2
        ly = (start[1] + end[1]) // 2 - 28
        draw.text((lx - 42, ly), label, font=font(18, True), fill="#334155")
    draw.text((785, 680), "Camadas internas: API -> Application -> Domain -> Infrastructure", font=font(23, True), fill="#1e3a8a")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def make_data_model_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1200), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((55, 35), "Modelo de dados do AutoDownload", font=font(36, True), fill="#0f172a")

    entities = {
        "users": ((70, 145, 460, 335), ["Id", "Name", "Email", "PasswordHash", "CreatedAt", "UpdatedAt"]),
        "operators": ((1280, 145, 1720, 335), ["Id", "Code", "Name", "ServiceType", "PortalBaseUrl", "IsActive"]),
        "accounts": (
            (560, 145, 1180, 455),
            [
                "Id",
                "UserId",
                "OperatorId",
                "PortalLogin",
                "CustomerIdentifier",
                "Status",
                "LastRunAt",
                "NextRunAt",
                "IsScheduleEnabled",
                "ScheduleDayOfMonth",
                "ScheduleTime",
            ],
        ),
        "bills": ((70, 640, 520, 940), ["Id", "UserId", "AccountId", "OperatorId", "Reference", "DueDate", "Amount", "FileName", "StoragePath", "DownloadedAt", "Status"]),
        "automation_runs": ((650, 640, 1120, 940), ["Id", "UserId", "AccountId", "OperatorId", "StartedAt", "FinishedAt", "Status", "Message", "FileName"]),
        "notifications": ((1260, 640, 1720, 860), ["Id", "UserId", "Type", "Text", "CreatedAt", "ReadAt"]),
    }
    for name, (box, fields) in entities.items():
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=16, fill="#f8fafc", outline="#2563eb", width=3)
        draw.rectangle((x1, y1, x2, y1 + 48), fill="#1d4ed8")
        draw.text((x1 + 18, y1 + 10), name, font=font(24, True), fill="#ffffff")
        y = y1 + 62
        for field in fields:
            draw.text((x1 + 18, y), field, font=font(19), fill="#0f172a")
            y += 24

    relationships = [
        ("users", "accounts", "1", "N"),
        ("operators", "accounts", "1", "N"),
        ("accounts", "bills", "1", "N"),
        ("accounts", "automation_runs", "1", "N"),
        ("users", "notifications", "1", "N"),
    ]
    centers = {
        k: ((v[0][0] + v[0][2]) // 2, (v[0][1] + v[0][3]) // 2)
        for k, v in entities.items()
    }
    for a, b, ca, cb in relationships:
        start, end = centers[a], centers[b]
        line(draw, start, end, "#475569", 3)
        draw.text((start[0] + 10, start[1] + 8), ca, font=font(20, True), fill="#0f172a")
        draw.text((end[0] + 10, end[1] + 8), cb, font=font(20, True), fill="#0f172a")

    draw.text((70, 1080), "Campos de agendamento mensal incluídos em accounts: IsScheduleEnabled, ScheduleDayOfMonth e ScheduleTime.", font=font(22, True), fill="#1e3a8a")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def clear_paragraph(paragraph: Paragraph) -> None:
    paragraph.clear()


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_para(doc: Document, needle: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    return None


def para_index(doc: Document, paragraph: Paragraph) -> int:
    return next(i for i, candidate in enumerate(doc.paragraphs) if candidate._element is paragraph._element)


def set_text(paragraph: Paragraph, text: str, bold: bool = False, size: int | None = None, align=None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if align is not None:
        paragraph.alignment = align


def replace_text_preserve_runs(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        if run._element.xpath(".//pic:pic"):
            continue
        txt = run.text
        for old, new in replacements.items():
            txt = txt.replace(old, new)
        run.text = txt


def set_image_paragraph(paragraph: Paragraph, image_path: Path, width_cm: float = 16.0) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    paragraph.add_run("\nFonte: Elaborado pelo autor (2026).")


def set_image_caption_in_same_paragraph(paragraph: Paragraph, caption: str) -> None:
    done = False
    for run in paragraph.runs:
        if run._element.xpath(".//pic:pic"):
            continue
        if not done:
            run.text = f"\n{caption}\nFonte: Elaborado pelo autor (2026)."
            done = True
        else:
            run.text = ""


def style_paragraph(paragraph: Paragraph, size: int = 12, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold


def style_table(table, widths_cm: list[float] | None = None) -> None:
    for style_name in ("Table Grid", "Tabela com grade"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = "Arial"
                    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
                    run.font.size = Pt(9 if widths_cm and len(widths_cm) > 4 else 10)
                    if row_idx == 0:
                        run.font.bold = True
            if widths_cm and col_idx < len(widths_cm):
                cell.width = Cm(widths_cm[col_idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            shade = tc_pr.find(qn("w:shd"))
            if row_idx == 0:
                if shade is None:
                    shade = OxmlElement("w:shd")
                    tc_pr.append(shade)
                shade.set(qn("w:fill"), "D9EAF7")


def fill_table(table, rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    expected_cols = len(rows[0])
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    needs_rebuild = any(len(table.rows[i].cells) < expected_cols for i in range(len(rows)))
    if needs_rebuild:
        new_table = table._parent.add_table(rows=len(rows), cols=expected_cols, width=Cm(16))
        table._tbl.addprevious(new_table._tbl)
        table._tbl.getparent().remove(table._tbl)
        table = new_table
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = value
    style_table(table, widths_cm)


def add_table_before(anchor: Paragraph, rows: list[list[str]], widths_cm: list[float] | None = None):
    doc = anchor._parent
    table = doc.add_table(rows=len(rows), cols=len(rows[0]), width=Cm(16))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    style_table(table, widths_cm)
    anchor._p.addprevious(table._tbl)
    return table


def add_para_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    use_case = IMG_DIR / "figura_1_casos_uso.png"
    architecture = IMG_DIR / "figura_2_arquitetura.png"
    data_model = IMG_DIR / "figura_3_modelo_dados.png"
    make_use_case_diagram(use_case)
    make_architecture_diagram(architecture)
    make_data_model_diagram(data_model)

    doc = Document(str(SRC))
    original_tables = list(doc.tables)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    set_text(
        doc.paragraphs[0],
        "AUTODOWNLOAD: UMA APLICAÇÃO WEB PARA AUTOMAÇÃO E GESTÃO DE BOLETOS RECORRENTES",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    if len(doc.paragraphs) > 1 and "Uma Aplicação Web" in doc.paragraphs[1].text:
        delete_paragraph(doc.paragraphs[1])

    replacements = {
        "Tecnologias Digitais Aplicada a Educação": "Tecnologias Digitais Aplicadas à Educação",
        "Tecnologias Digitais Aplicadas a Educação": "Tecnologias Digitais Aplicadas à Educação",
        "Analise e Desenvolvimento de Sistemas": "Análise e Desenvolvimento de Sistemas",
        "Aplicação web. NET. Next.js. PostgreSQL. Selenium.": ".NET. PostgreSQL. Selenium.",
        "NET.": ".NET.",
        "internet, energia, água, telefone ou outros serviços recorrentes": "internet, telecomunicações e serviços recorrentes similares",
        "contas de internet, energia, água, telefone ou outros serviços recorrentes": "contas de internet e telecomunicações, além de outros serviços recorrentes tratados como possibilidade futura",
        "serviços financeiros, de telecomunicações e de utilidades públicas": "serviços financeiros e de telecomunicações",
        "contribui para reduzir tarefas repetitivas": "apresenta potencial para reduzir tarefas repetitivas",
        "os resultados indicam": "os resultados observados indicam, de forma preliminar",
    }
    for paragraph in doc.paragraphs:
        replace_text_preserve_runs(paragraph, replacements)

    kw = find_para(doc, "Palavras-chave:")
    if kw:
        set_text(kw, "Palavras-chave: Automação web. Boletos recorrentes. .NET. PostgreSQL. Selenium.", size=12)

    resumo = find_para(doc, "O presente Trabalho de Conclusão de Curso apresenta")
    if resumo:
        set_text(
            resumo,
            "O presente Trabalho de Conclusão de Curso apresenta o desenvolvimento do AutoDownload, "
            "uma aplicação web voltada à automação do acesso a portais de operadoras de internet e "
            "telecomunicações e à gestão de boletos recorrentes. O sistema foi desenvolvido com frontend "
            "em Next.js e React, backend em C# com .NET 10, persistência em PostgreSQL por meio do Entity "
            "Framework Core e automações com Selenium WebDriver. A aplicação contempla autenticação JWT, "
            "cadastro de contas, execução manual de automações, agendamento mensal, histórico, notificações "
            "e registro de boletos. A validação incluiu testes unitários no backend e uso controlado da "
            "aplicação, indicando que a solução apresenta potencial para reduzir tarefas repetitivas, "
            "desde que consideradas as limitações de portais externos e de ambientes gratuitos de hospedagem.",
            size=12,
        )

    correlatos = find_para(doc, "A análise de trabalhos correlatos permite")
    if correlatos:
        set_text(
            correlatos,
            "A análise de trabalhos correlatos considerou soluções concretas usadas no contexto do projeto: "
            "os portais Minha Vero e RMS Telecom, aplicativos bancários que permitem consulta de boletos "
            "previamente cadastrados e o repositório AutoBot, utilizado como referência técnica inicial. "
            "Os critérios de comparação foram centralização de contas, automação do download, histórico, "
            "notificações e limitações operacionais. Essa comparação delimita a contribuição do AutoDownload, "
            "que busca integrar em uma única interface cadastro de contas, automação por estratégia de "
            "operadora, registro em banco de dados, histórico e notificações.",
            size=12,
        )

    fill_table(
        original_tables[0],
        [
            ["Solução analisada", "Centraliza contas", "Automatiza download", "Histórico", "Notificações", "Limitações"],
            ["Portal Minha Vero", "Não", "Parcial", "Parcial", "Parcial", "Atende apenas a própria operadora e exige acesso separado ao portal."],
            ["Portal RMS Telecom", "Não", "Parcial", "Parcial", "Parcial", "Atende apenas a própria operadora e depende da navegação manual do usuário."],
            ["Aplicativos bancários com boletos cadastrados", "Parcial", "Não", "Parcial", "Sim", "Dependem de boletos já recebidos, cadastrados ou pagos pelo usuário."],
            ["AutoBot", "Não", "Sim, para fluxo específico", "Limitado", "Não integrado", "Automação isolada, sem dashboard, persistência estruturada e gestão multioperadora."],
            ["AutoDownload", "Sim", "Sim, por estratégia de operadora", "Sim", "Sim", "Depende da estabilidade dos portais externos e de evidências contínuas de execução real."],
        ],
        [3.0, 2.0, 3.0, 2.0, 2.0, 4.0],
    )

    methodology = find_para(doc, "A validação ocorreu por meio")
    if methodology:
        set_text(
            methodology,
            "A validação ocorreu por meio de build da solução, execução da suíte de testes unitários do backend "
            "com xUnit, verificação do endpoint de saúde da API, uso manual da aplicação e execuções controladas "
            "das automações. A suíte foi executada com o comando "
            f"{TEST_RESULT} As automações reais e o agendamento mensal são registrados como validações parciais "
            "quando não acompanhados de log, captura, arquivo obtido e registro de banco no artigo.",
            size=12,
        )
        extra = add_para_after(methodology, "A persona utilizada no apêndice foi construída de forma sintética, a partir da observação do fluxo manual de obtenção de boletos e das necessidades identificadas durante o desenvolvimento do protótipo.")
        style_paragraph(extra)

    dev_heading = find_para(doc, "4 DESENVOLVIMENTO")
    if dev_heading:
        cap = dev_heading.insert_paragraph_before("Quadro 2 - Procedimentos metodológicos adotados")
        style_paragraph(cap, bold=True)
        add_table_before(
            dev_heading,
            [
                ["Fonte", "Período", "Instrumento", "Procedimento", "Produto gerado"],
                ["Observação do fluxo manual de boletos", "2026/1", "Roteiro de etapas", "Mapeamento do acesso a portais, autenticação, consulta e download", "Problema, persona sintética e requisitos iniciais"],
                ["Documentos do projeto e pareceres", "2026/1", "Análise documental", "Leitura de proposta, orientações e retornos de revisão", "Escopo, limitações e ajustes acadêmicos"],
                ["Repositório e aplicação local", "2026/1", "Build, testes e uso manual", "Execução da API, frontend, banco e testes unitários", "Evidências técnicas e quadro de validação"],
                ["Portais externos", "2026/1", "Execução controlada", "Teste pontual das estratégias de automação quando disponíveis", "Classificação parcial quando sem evidência completa anexada"],
            ],
            [3.2, 2.0, 3.0, 4.2, 4.0],
        )
        source = dev_heading.insert_paragraph_before("Fonte: Elaborado pelo autor (2026).")
        style_paragraph(source)

    # Captions and diagrams in the body.
    caption_map = {
        "Figura - 1": "Figura 1 - Diagrama de casos de uso do AutoDownload",
        "Figura - 2": "Figura 2 - Arquitetura geral do AutoDownload",
        "Figura - 3": "Figura 3 - Modelo de dados do AutoDownload",
        "Quadro - 2": "Quadro 3 - Requisitos funcionais",
        "Quadro - 3": "Quadro 4 - Requisitos não funcionais",
        "Quadro - 4": "Quadro 5 - Principais entidades do banco de dados",
        "Quadro - 5": "Quadro 6 - Validação funcional",
        "Quadro - 6": "Quadro 7 - Baseline técnico e contribuições do AutoDownload",
    }
    for paragraph in doc.paragraphs:
        txt = paragraph.text.strip()
        for old, new in caption_map.items():
            if txt.startswith(old):
                set_text(paragraph, new, bold=True, size=12)

    for caption, image_path in [
        ("Figura 1 - Diagrama de casos de uso do AutoDownload", use_case),
        ("Figura 2 - Arquitetura geral do AutoDownload", architecture),
        ("Figura 3 - Modelo de dados do AutoDownload", data_model),
    ]:
        cap = find_para(doc, caption)
        if cap:
            image_p = None
            paragraphs = doc.paragraphs
            idx = para_index(doc, cap)
            for candidate in paragraphs[idx + 1 : idx + 4]:
                if candidate._element.xpath(".//pic:pic"):
                    image_p = candidate
                    break
            if image_p:
                set_image_paragraph(image_p, image_path, 16.0)

    # Screenshot captions in existing image paragraphs.
    screenshot_captions = {
        "Figura - 01:": "Figura 4 - Tela de login",
        "Figura - 02:": "Figura 5 - Tela inicial da aplicação",
        "Figura - 03:": "Figura 6 - Tela para adicionar uma conta",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if paragraph._element.xpath(".//pic:pic"):
            for old, new in screenshot_captions.items():
                if old in text:
                    set_image_caption_in_same_paragraph(paragraph, new)

    fill_table(
        original_tables[1],
        [
            ["Código", "Requisito", "Situação"],
            ["RF01", "Permitir cadastro de usuário", "Implementado"],
            ["RF02", "Permitir login de usuário", "Implementado"],
            ["RF03", "Permitir cadastro de contas vinculadas a operadoras", "Implementado"],
            ["RF04", "Permitir edição e remoção de contas", "Implementado"],
            ["RF05", "Executar automação manual de download", "Implementado"],
            ["RF06", "Permitir consulta e download dos boletos registrados", "Implementado"],
            ["RF07", "Registrar histórico de execuções", "Implementado"],
            ["RF08", "Exibir notificações de sucesso, falha ou alterações", "Implementado"],
            ["RF09", "Permitir configuração de agendamento mensal", "Implementado parcialmente"],
            ["RF10", "Recuperar senha por e-mail com token", "Não implementado na versão acadêmica"],
        ],
        [2.2, 10.0, 4.0],
    )
    fill_table(
        original_tables[2],
        [
            ["Código", "Requisito", "Tratamento no projeto"],
            ["RNF01", "Segurança de acesso", "Autenticação JWT e rotas autenticadas na API"],
            ["RNF02", "Proteção de senhas de usuários", "Armazenamento por hash, sem persistência de senha em texto puro"],
            ["RNF03", "Proteção de credenciais externas", "Criptografia das senhas de portais com chave configurada por variável de ambiente"],
            ["RNF04", "Manutenibilidade", "Arquitetura em camadas, injeção de dependência, Repository e Strategy"],
            ["RNF05", "Persistência", "PostgreSQL com Entity Framework Core e migrations"],
            ["RNF06", "Rastreabilidade", "Histórico, notificações e registros de execução"],
            ["RNF07", "Usabilidade", "Interface web com dashboard, filtros e modo claro/escuro"],
            ["RNF08", "Disponibilidade", "Frontend, backend e banco preparados para deploy em serviços separados"],
        ],
        [2.2, 5.0, 8.8],
    )
    fill_table(
        original_tables[3],
        [
            ["Entidade", "Finalidade"],
            ["users", "Armazena usuários da aplicação, dados de autenticação e datas de criação/atualização."],
            ["operators", "Armazena operadoras disponíveis, código, nome, tipo de serviço e URL base."],
            ["accounts", "Armazena contas cadastradas, vínculo com usuário/operadora, login do portal, identificador, status, última e próxima execução, IsScheduleEnabled, ScheduleDayOfMonth e ScheduleTime."],
            ["bills", "Armazena boletos baixados, referência, vencimento, valor, arquivo, caminho de armazenamento e status."],
            ["automation_runs", "Registra execuções de automação, início, fim, status, mensagem e arquivo gerado."],
            ["notifications", "Armazena notificações do usuário, tipo, texto, data de criação e data de leitura."],
        ],
        [4.0, 11.8],
    )
    fill_table(
        original_tables[4],
        [
            ["Requisito", "Cenário testado", "Resultado esperado", "Resultado obtido", "Situação"],
            ["Login", "Credenciais válidas na aplicação", "Acesso liberado e token gerado", "Acesso realizado em ambiente local e publicado", "Aprovado"],
            ["Cadastro de conta", "Usuário informa operadora, login, senha e identificador", "Conta persistida no banco", "Fluxo implementado e validado manualmente", "Aprovado"],
            ["Edição e remoção de conta", "Usuário altera ou exclui conta cadastrada", "Dados atualizados ou removidos", "Operação refletida na interface e no banco", "Aprovado"],
            ["Automação manual", "Usuário solicita download de boleto", "Execução registrada e boleto salvo quando disponível", "Validada pontualmente em ambiente local; em hospedagem gratuita depende de recursos do container", "Parcial"],
            ["Automação real Vero/RMS", "Acesso ao portal externo", "Boleto disponível baixado e registrado", "Sem evidência completa anexada de data, log, arquivo, histórico e banco", "Parcial"],
            ["Agendamento mensal", "Usuário define dia e horário", "Execução automática recorrente", "Configuração exibida; execução recorrente ainda exige evidência completa", "Parcial"],
            ["Histórico e notificações", "Execuções geram registros para o usuário autenticado", "Histórico e notificações filtrados por usuário", "Fluxo implementado e revisado para não exibir dados mockados", "Aprovado"],
            ["Testes automatizados backend", "Execução de dotnet test em 25/06/2026", "Suíte sem falhas", "20 testes aprovados, 0 falhas, 0 ignorados, duração 182 ms", "Aprovado"],
        ],
        [2.8, 3.8, 3.8, 4.5, 2.0],
    )

    # Body technology paragraph with citations for references already listed.
    frontend_para = find_para(doc, "No frontend, Next.js e React foram utilizados")
    if frontend_para:
        set_text(
            frontend_para,
            "No frontend, Next.js e React foram utilizados para construir a interface da aplicação. "
            "A comunicação com o backend ocorre por chamadas HTTP para a API REST. A publicação do frontend "
            "foi preparada para a Vercel, plataforma compatível com aplicações Next.js e pipelines de deploy "
            "por repositório Git (VERCEL, 2026).",
            size=12,
        )
    backend_para = find_para(doc, "No backend, .NET 10 foi utilizado")
    if backend_para:
        set_text(
            backend_para,
            "No backend, .NET 10 foi utilizado para criar uma API REST. A autenticação foi implementada com JWT, "
            "permitindo identificar o usuário autenticado e filtrar informações por proprietário. As senhas dos "
            "usuários são armazenadas por hash, enquanto credenciais dos portais são protegidas por criptografia. "
            "O acesso ao PostgreSQL foi implementado com Entity Framework Core, conforme a documentação oficial "
            "da Microsoft para ASP.NET Core e EF Core (MICROSOFT, 2026a; MICROSOFT, 2026b).",
            size=12,
        )

    # Conclusion moderation.
    for paragraph in doc.paragraphs:
        if "Como limitação" in paragraph.text and "portais externos" in paragraph.text:
            set_text(
                paragraph,
                "Como limitação, destaca-se que automações dependentes de portais externos podem falhar diante "
                "de mudanças na interface, indisponibilidade, validações adicionais ou restrições de containers "
                "em hospedagem gratuita. Assim, os resultados devem ser interpretados como evidência técnica "
                "preliminar do potencial da solução, e não como comprovação estatística de redução de tempo ou "
                "esforço do usuário.",
                size=12,
            )

    # References.
    refs = [
        "AUTOBOT. Repositório AutoBot. GitHub, 2026. Disponível em: https://github.com/Edyeex/AutoBot. Acesso em: 25 jun. 2026.",
        "AUTODOWNLOAD. Repositório AutomacaoContas. GitHub, 2026. Commit 6cc09724a6c5e1ba0c85abe6ea7c0b5e6ed88858. Disponível em: https://github.com/Edyeex/AutomacaoContas. Acesso em: 25 jun. 2026.",
        "BRASIL. Lei n. 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais (LGPD). Brasília, DF: Presidência da República, 2018. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709.htm. Acesso em: 25 jun. 2026.",
        "BROWN, Tim. Design thinking. Harvard Business Review, Brighton, jun. 2008. Disponível em: https://hbr.org/2008/06/design-thinking. Acesso em: 25 jun. 2026.",
        "GIL, Antonio Carlos. Métodos e técnicas de pesquisa social. 7. ed. São Paulo: Atlas, 2017.",
        "MICROSOFT. ASP.NET Core documentation. 2026a. Disponível em: https://learn.microsoft.com/aspnet/core/. Acesso em: 25 jun. 2026.",
        "MICROSOFT. Entity Framework Core documentation. 2026b. Disponível em: https://learn.microsoft.com/ef/core/. Acesso em: 25 jun. 2026.",
        "NORMAN, Donald A. O design do dia a dia. Rio de Janeiro: Rocco, 2006.",
        "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de software: uma abordagem profissional. 8. ed. Porto Alegre: AMGH, 2016.",
        "RMS TELECOM. Portal do assinante RMS Telecom. 2026. Disponível em: https://rmstelecom.net/. Acesso em: 25 jun. 2026.",
        "SELENIUM. Selenium WebDriver documentation. 2026. Disponível em: https://www.selenium.dev/documentation/webdriver/. Acesso em: 25 jun. 2026.",
        "SOMMERVILLE, Ian. Engenharia de software. 10. ed. São Paulo: Pearson, 2019.",
        "VERCEL. Next.js on Vercel. 2026. Disponível em: https://vercel.com/docs/frameworks/nextjs. Acesso em: 25 jun. 2026.",
        "VERO INTERNET. Minha Vero. 2026. Disponível em: https://verointernet.com.br/minhavero/login. Acesso em: 25 jun. 2026.",
    ]
    ref_heading = find_para(doc, "REFERÊNCIAS")
    if ref_heading:
        paragraphs = doc.paragraphs
        start = para_index(doc, ref_heading) + 1
        end = next((i for i in range(start, len(paragraphs)) if paragraphs[i].text.strip().startswith("APÊNDICES")), len(paragraphs))
        for paragraph in list(paragraphs[start:end]):
            delete_paragraph(paragraph)
        anchor = ref_heading
        for item in reversed(refs):
            p = add_para_after(anchor, item)
            style_paragraph(p, size=12)

    ap = find_para(doc, "APÊNDICES SUGERIDOS")
    if ap:
        set_text(ap, "APÊNDICES", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        ap = find_para(doc, "APÊNDICES")
        if ap:
            set_text(ap, "APÊNDICES", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Appendix A: explicit synthetic persona.
    for table in doc.tables:
        if len(table.rows) >= 2 and table.cell(0, 0).text.strip() == "Campo" and "Persona" in table.cell(1, 0).text:
            table.cell(1, 1).text = "Usuário doméstico responsável pelo pagamento de contas. Persona sintética criada a partir do fluxo observado no projeto."
            style_table(table, [4.0, 11.5])
        if len(table.rows) >= 1 and table.cell(0, 0).text.strip().lower() == "etapa":
            table.cell(0, 0).text = "Etapa"
            style_table(table, [3.0, 4.0, 4.2, 4.5])

    # Appendix B: replace duplicate use-case diagram with traceability matrix.
    app_b = find_para(doc, "APÊNDICE B -")
    app_c = find_para(doc, "APÊNDICE C -")
    if app_b and app_c:
        set_text(app_b, "APÊNDICE B - Matriz de rastreabilidade requisito-teste.", bold=True, size=12)
        start = para_index(doc, app_b) + 1
        end = para_index(doc, app_c)
        for paragraph in list(doc.paragraphs[start:end]):
            delete_paragraph(paragraph)
        matrix = [
            ["Requisito", "Implementação/evidência", "Situação"],
            ["RF01-RF02", "Fluxos de cadastro/login, autenticação JWT e testes de serviços", "Implementado"],
            ["RF03-RF04", "Cadastro, edição e remoção de contas com persistência no banco", "Implementado"],
            ["RF05-RF07", "Execução de automação, boletos e histórico", "Implementado com evidência parcial para portais reais"],
            ["RF08", "Notificações por usuário para sucesso, falha e alterações", "Implementado"],
            ["RF09", "Configuração de dia e horário mensal por conta", "Implementado parcialmente; execução recorrente exige evidência final"],
            ["RF10", "Recuperação de senha por e-mail com token", "Não implementado na versão acadêmica"],
            ["RNF01-RNF08", "JWT, hash, criptografia, camadas, EF Core, PostgreSQL, rastreabilidade e deploy", "Implementado conforme escopo"],
        ]
        src = add_para_after(app_b, "Fonte: Elaborado pelo autor (2026).")
        style_paragraph(src)
        tbl = doc.add_table(rows=len(matrix), cols=len(matrix[0]))
        for r, row in enumerate(matrix):
            for c, value in enumerate(row):
                tbl.cell(r, c).text = value
        style_table(tbl, [3.0, 10.5, 3.2])
        src._p.addprevious(tbl._tbl)

    # Appendix C: replace duplicate architecture/data diagrams with evidence note.
    app_c = find_para(doc, "APÊNDICE C -")
    app_d = find_para(doc, "APÊNDICE D -")
    if app_c and app_d:
        set_text(app_c, "APÊNDICE C - Evidências de validação técnica.", bold=True, size=12)
        start = para_index(doc, app_c) + 1
        end = para_index(doc, app_d)
        for paragraph in list(doc.paragraphs[start:end]):
            delete_paragraph(paragraph)
        p1 = add_para_after(app_c, "Resultado registrado para a versão acadêmica: " + TEST_RESULT)
        style_paragraph(p1)
        rows = [
            ["Evidência", "Registro"],
            ["Build/testes unitários", "20 testes aprovados, 0 falhas, 0 ignorados, duração 182 ms"],
            ["API publicada", "Endpoint /api/health retornando status healthy em ambiente Render"],
            ["Frontend publicado", "Aplicação disponível em https://automacao-contas.vercel.app/"],
            ["Automação real/agendamento", "Manter como parcial até anexar log, arquivo, banco, histórico e captura do ciclo completo"],
        ]
        src = add_para_after(p1, "Fonte: Elaborado pelo autor (2026).")
        style_paragraph(src)
        tbl = doc.add_table(rows=len(rows), cols=2)
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                tbl.cell(r, c).text = value
        style_table(tbl, [4.2, 11.2])
        src._p.addprevious(tbl._tbl)

    app_d = find_para(doc, "APÊNDICE D -")
    if app_d:
        set_text(app_d, "APÊNDICE D - Capturas complementares da interface.", bold=True, size=12)
        note = add_para_after(
            app_d,
            "As capturas deste apêndice documentam a interface. Para a entrega final, recomenda-se substituir ou complementar as imagens por evidências populadas e anonimizadas, mostrando conta cadastrada, agendamento ativo, execução concluída, boleto registrado, notificação e histórico.",
        )
        style_paragraph(note)

    # Appendix D captions near screenshots.
    app_d_caption_replacements = {
        "Figura - D1:": "Figura D1 - Tela de login",
        "Figura - D2:": "Figura D2 - Tela principal",
        "Figura - D3:": "Figura D3 - Tela de contas cadastradas e cadastro de contas",
        "Figura - D4:": "Figura D4 - Tela de notificações",
        "Figura - D5:": "Figura D5 - Tela de histórico de downloads",
    }
    for paragraph in doc.paragraphs:
        txt = paragraph.text
        for old, new in app_d_caption_replacements.items():
            if old in txt:
                set_text(paragraph, new, bold=True, size=12)
        if paragraph._element.xpath(".//pic:pic") and not paragraph.text.strip():
            # Keep image-only paragraphs untouched.
            pass

    app_e = find_para(doc, "APÊNDICE E -")
    if app_e:
        for paragraph in doc.paragraphs[para_index(doc, app_e) + 1 :]:
            txt = paragraph.text.strip()
            if txt.startswith("Repositório do projeto:"):
                set_text(paragraph, "Repositório do projeto: https://github.com/Edyeex/AutomacaoContas", size=12)
            elif txt.startswith("Versão acadêmica/tag:"):
                set_text(paragraph, "Versão acadêmica/tag: não publicada nesta versão; identificação realizada pelo hash do commit.", size=12)
            elif txt.startswith("Hash do commit:"):
                set_text(paragraph, "Hash do commit: 6cc09724a6c5e1ba0c85abe6ea7c0b5e6ed88858", size=12)
            elif txt.startswith("Data de validação:"):
                set_text(paragraph, "Data de validação: 25/06/2026", size=12)

    # General heading/caption styling.
    for paragraph in doc.paragraphs:
        stripped = paragraph.text.strip()
        if re.match(r"^\d(\.\d)?\s", stripped) or stripped in {"RESUMO", "REFERÊNCIAS", "APÊNDICES"} or stripped.startswith("APÊNDICE"):
            style_paragraph(paragraph, 12, bold=True)
        elif stripped.startswith(("Figura ", "Quadro ", "Fonte:")):
            style_paragraph(paragraph, 12)
        else:
            style_paragraph(paragraph, 12)

    doc.core_properties.author = "Éder Casagranda"
    doc.core_properties.comments = "Arquivo ajustado conforme Parecer de Orientação Acadêmica - versão 5."
    doc.save(str(OUT))

    LOG.write_text(
        "\n".join(
            [
                f"Arquivo gerado: {OUT}",
                "Original preservado: C:\\Users\\android\\Downloads\\AUTODOWNLOAD.docx",
                "Ajustes principais:",
                "- Margens A4 ajustadas para 3 cm superior/esquerda e 2 cm inferior/direita.",
                "- Título consolidado em caixa alta, Arial 14 e negrito.",
                "- Resumo e conclusão reescritos com linguagem cautelosa.",
                "- Escopo textual limitado a internet/telecomunicações.",
                "- Trabalhos correlatos reformulados com soluções concretas.",
                "- Quadro metodológico inserido.",
                "- Diagramas de casos de uso, arquitetura e modelo de dados substituídos.",
                "- Requisitos funcionais e não funcionais corrigidos.",
                "- Modelo de dados atualizado com campos de agendamento.",
                "- Validação corrigida com resultado real de dotnet test.",
                "- Recuperação de senha marcada como não implementada.",
                "- Tag acadêmica não publicada removida do apêndice E.",
                "- Apêndices B e C substituídos por matriz de rastreabilidade e evidências técnicas.",
                "- Referências revisadas e complementadas.",
            ]
        ),
        encoding="utf-8",
    )
    print(OUT)


if __name__ == "__main__":
    main()
